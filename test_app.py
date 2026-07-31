import sys
import os
import json
import docx
import pandas as pd
from fastapi.testclient import TestClient

# Add parent directory to path so we can import app modules
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from app.main import app

client = TestClient(app)

def create_mock_docx(filepath: str, content: str):
    """
    Creates a mock Word document for testing.
    """
    doc = docx.Document()
    doc.add_heading("Mock Candidate Profile", 0)
    doc.add_paragraph(content)
    doc.save(filepath)

def create_mock_excel(filepath: str):
    """
    Creates a mock Excel sheet with candidate names and resumes.
    """
    data = {
        "Candidate Name": [
            "Emma Watson - Frontend Engineer",
            "Liam Neeson - DevOps Architect"
        ],
        "Resume Content": [
            "Emma Watson. Senior React Developer with 5 years experience in TypeScript, React Redux, Next.js, and CSS layout optimization.",
            "Liam Neeson. DevOps Architect. Skills include AWS, Terraform, Docker, Kubernetes orchestration, CI/CD, and server provisioning."
        ]
    }
    df = pd.DataFrame(data)
    df.to_excel(filepath, index=False)

def run_tests():
    print("=" * 70)
    print("Running Extended API Integration Tests (Word + Excel Bulk)...")
    print("=" * 70)

    # 1. Clear database
    print("\n[Step 1] Clearing vector store database...")
    response = client.delete("/resumes")
    assert response.status_code == 200
    print(f"Response: {response.json()}")

    # 2. Test Word Document Upload
    print("\n[Step 2] Testing individual Word (.docx) resume upload...")
    mock_docx_path = "data/mock_candidate_react.docx"
    os.makedirs("data", exist_ok=True)
    create_mock_docx(
        mock_docx_path,
        "John Doe. Backend Developer. Experienced in Python, FastAPI, Docker, and PostgreSQL databases."
    )

    with open(mock_docx_path, "rb") as f:
        response = client.post(
            "/upload-resume",
            files={"file": (os.path.basename(mock_docx_path), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
    
    assert response.status_code == 200, f"Word upload failed: {response.text}"
    res_data = response.json()
    assert res_data["is_bulk"] is False
    assert res_data["count"] == 1
    assert "mock_candidate_react.docx" in res_data["filename"]
    print(f"Success: Individual Word file parsed & indexed -> Candidate File: {res_data['filename']}")

    # Clean up mock file
    if os.path.exists(mock_docx_path):
        os.remove(mock_docx_path)

    # 3. Test Bulk Excel Upload
    print("\n[Step 3] Testing bulk Excel (.xlsx) resumes upload...")
    mock_excel_path = "data/mock_bulk_candidates.xlsx"
    create_mock_excel(mock_excel_path)

    with open(mock_excel_path, "rb") as f:
        response = client.post(
            "/upload-resume",
            files={"file": (os.path.basename(mock_excel_path), f, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
        )

    assert response.status_code == 200, f"Excel bulk upload failed: {response.text}"
    res_bulk_data = response.json()
    assert res_bulk_data["is_bulk"] is True
    assert res_bulk_data["count"] == 2
    print(f"Success: Excel sheet parsed -> {res_bulk_data['message']}")

    # Clean up mock file
    if os.path.exists(mock_excel_path):
        os.remove(mock_excel_path)

    # 4. Verify Total Candidates Indexed
    print("\n[Step 4] Verifying total indexed candidates count...")
    response = client.get("/resumes")
    assert response.status_code == 200
    resumes_list = response.json()
    print(f"Total resumes currently in vector store: {len(resumes_list)}")
    assert len(resumes_list) == 3, f"Expected 3 candidates, got {len(resumes_list)}"

    # 5. Test Matching Query (React/Frontend Candidate search)
    print("\n[Step 5] Performing S-BERT match query for React developer...")
    response = client.post(
        "/match",
        json={"job_description": "We are seeking a senior frontend UI developer with expertise in React, Next.js, and TypeScript.", "top_k": 2}
    )
    assert response.status_code == 200
    match_data = response.json()
    
    top_match = match_data["matches"][0]
    print(f"Top Match: {top_match['filename']} (Score: {top_match['score']:.4f})")
    assert "Emma Watson" in top_match["filename"], "Expected Emma Watson (Frontend) to be the top match!"

    # 6. Test RAG Analysis on the bulk-uploaded candidate
    print("\n[Step 6] Running RAG Fit Analysis on bulk candidate...")
    rag_request = {
        "resume_id": top_match["resume_id"],
        "job_description": "Seeking a frontend web developer with React and TypeScript skills."
    }
    response = client.post("/rag-analyze", json=rag_request)
    assert response.status_code == 200
    rag_data = response.json()
    print("RAG AI fit evaluation generated successfully:")
    print("-" * 60)
    print(rag_data["analysis"])
    print("-" * 60)

    print("\nAll integration tests (Word and Excel Bulk Upload) passed successfully! [OK]")

if __name__ == "__main__":
    run_tests()
