import os
import io
import docx
from pypdf import PdfReader

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts plain text from PDF file bytes.
    """
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts plain text from Word DOCX file bytes.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text).strip()
    except Exception as e:
        raise ValueError(f"Error parsing Word DOCX file: {str(e)}")

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Extracts text from bytes based on file extension.
    """
    ext = os.path.splitext(filename.lower())[1]
    if ext == '.pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext == '.docx':
        return extract_text_from_docx(file_bytes)
    elif ext in ['.txt', '.md']:
        try:
            return file_bytes.decode('utf-8', errors='ignore').strip()
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF, DOCX and TXT/MD files are supported.")
